from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
MD_PATH = BASE_DIR / "Bao_Cao_Do_An_Tuan3.md"
OUT_PATH = BASE_DIR / "25521787_Bao_Cao_Do_An_Tuan3_IT003_Q21_CTTN_Final.docx"


ACCENT = RGBColor(0x00, 0x66, 0xA6)
ACCENT_DARK = RGBColor(0x12, 0x2E, 0x4A)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color


def set_cell_padding(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, ACCENT_DARK),
        ("Heading 2", 14, ACCENT),
        ("Heading 3", 12.5, ACCENT_DARK),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Cm(0.45)
    code.paragraph_format.right_indent = Cm(0.2)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO ĐỒ ÁN MÔN HỌC")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(20)
    r.font.color.rgb = ACCENT_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Đồ án phát triển ứng dụng")
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    rows = [
        ("Lớp", "IT003.Q21.CTTN"),
        ("Mã sinh viên", "25521787"),
        ("Họ và tên", "Võ Quốc Thịnh"),
        ("Tên đề tài", "Game Pacman"),
        ("Báo cáo", "Tuần 3"),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.color.rgb = ACCENT_DARK
        r = p.add_run(value)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CÁC NỘI DUNG CẦN BÁO CÁO")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.font.color.rgb = ACCENT


def add_rich_text(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xA0, 0x24, 0x24)
        else:
            paragraph.add_run(part)


def add_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        p.style = "Code Block"
        p.paragraph_format.space_after = Pt(1)
        shade_paragraph(p, "F5F7FA")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def build_docx():
    md = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    in_code = False
    code_lines = []
    skip_intro_until_rule = True

    for raw in md:
        line = raw.rstrip()
        if skip_intro_until_rule:
            if line.strip() == "---":
                skip_intro_until_rule = False
            continue

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph(text, style="Heading 1")
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph(text, style="Heading 2")
            continue

        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph(style="Heading 3")
            add_rich_text(p, line)
            continue

        if line.startswith("- "):
            content = line[2:].strip()
            if content.startswith("**") and content.endswith("**") and ":" not in content:
                p = doc.add_paragraph(style="Heading 3")
                add_rich_text(p, content)
            else:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.45)
                p.paragraph_format.space_after = Pt(3)
                add_rich_text(p, content)
            continue

        if re.match(r"^\s+- ", line):
            content = line.strip()[2:].strip()
            p = doc.add_paragraph(style="List Bullet 2")
            p.paragraph_format.left_indent = Cm(0.95)
            p.paragraph_format.space_after = Pt(2)
            add_rich_text(p, content)
            continue

        p = doc.add_paragraph()
        add_rich_text(p, line)

    if in_code:
        add_code_block(doc, code_lines)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Báo cáo đồ án tuần 3 - Game Pacman")
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_docx()
    print(OUT_PATH)
